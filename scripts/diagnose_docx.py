#!/usr/bin/env python3
"""
DOCX Diagnostic Script

Analyzes DOCX files to understand their structure and why the parser
might not be extracting issues correctly.

Usage:
    python3 scripts/diagnose_docx.py path/to/audit.docx
    python3 scripts/diagnose_docx.py path/to/directory/  # Analyzes first 3 files
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def analyze_paragraph_formatting(paragraph):
    """Analyze a paragraph's formatting to detect bullets."""
    info = {
        'text': paragraph.text.strip(),
        'style': paragraph.style.name,
        'is_list_item': False,
        'bullet_char': None,
        'first_chars': None
    }

    # Check if it's a list paragraph
    if paragraph._p.pPr is not None:
        numPr = paragraph._p.pPr.numPr
        if numPr is not None:
            info['is_list_item'] = True

    # Check first character
    if paragraph.text:
        info['first_chars'] = paragraph.text[:3]
        if paragraph.text[0] in ['•', '-', '*']:
            info['bullet_char'] = paragraph.text[0]

    return info


def analyze_document(docx_path: str, detailed: bool = False):
    """Analyze a DOCX file structure."""
    print(f"\n{'='*70}")
    print(f"📄 Analyzing: {os.path.basename(docx_path)}")
    print(f"{'='*70}")

    doc = Document(docx_path)

    # Track sections and issues
    sections_found = []
    potential_issues = []
    current_section = None

    # Analyze paragraph by paragraph
    for i, para in enumerate(doc.paragraphs[:100]):  # First 100 paragraphs
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name
        formatting = analyze_paragraph_formatting(para)

        # Detect sections (Heading 1 or Heading 2)
        if style in ['Heading 1', 'Heading 2']:
            if current_section:
                sections_found.append({
                    'title': current_section,
                    'issue_count': len([p for p in potential_issues if p['section'] == current_section])
                })
            current_section = text
            potential_issues.append({
                'type': 'SECTION_HEADER',
                'section': current_section,
                'text': text,
                'style': style
            })

        # Detect potential issues
        elif current_section and current_section not in ['Index', 'Audit Overview', 'Reference Websites']:
            # Check if it looks like an issue
            is_potential_issue = False
            reason = []

            # Method 1: Plain text bullet
            if formatting['bullet_char']:
                is_potential_issue = True
                reason.append(f"Plain bullet ({formatting['bullet_char']})")

            # Method 2: List item formatting
            if formatting['is_list_item']:
                is_potential_issue = True
                reason.append("List item formatting")

            # Method 3: Starts with common issue patterns
            if text.lower().startswith(('issue:', 'problem:', 'observation:')):
                is_potential_issue = True
                reason.append("Starts with issue keyword")

            # Method 4: Contains colon (title: description pattern)
            if ':' in text[:100] and len(text) > 20:
                is_potential_issue = True
                reason.append("Contains colon (title:description)")

            if is_potential_issue:
                potential_issues.append({
                    'type': 'POTENTIAL_ISSUE',
                    'section': current_section,
                    'text': text[:150],
                    'style': style,
                    'detection_method': ', '.join(reason),
                    'is_list_item': formatting['is_list_item'],
                    'bullet_char': formatting['bullet_char'],
                    'first_chars': formatting['first_chars']
                })

    # Save last section
    if current_section:
        sections_found.append({
            'title': current_section,
            'issue_count': len([p for p in potential_issues if p['section'] == current_section and p['type'] == 'POTENTIAL_ISSUE'])
        })

    # Print summary
    print(f"\n📊 SUMMARY:")
    print(f"  Total paragraphs analyzed: {len(doc.paragraphs[:100])}")
    print(f"  Sections found: {len(sections_found)}")
    print(f"  Potential issues detected: {len([p for p in potential_issues if p['type'] == 'POTENTIAL_ISSUE'])}")

    # Print sections with issue counts
    print(f"\n📋 SECTIONS DETECTED:")
    for section in sections_found:
        print(f"  • {section['title']}: {section['issue_count']} potential issues")

    # Show first few potential issues
    issues_list = [p for p in potential_issues if p['type'] == 'POTENTIAL_ISSUE']

    if issues_list:
        print(f"\n🔍 FIRST 5 POTENTIAL ISSUES:")
        for i, issue in enumerate(issues_list[:5], 1):
            print(f"\n  [{i}] Section: {issue['section']}")
            print(f"      Style: {issue['style']}")
            print(f"      Detection: {issue['detection_method']}")
            print(f"      Is list item: {issue['is_list_item']}")
            print(f"      Bullet char: {issue['bullet_char']}")
            print(f"      First chars: {repr(issue['first_chars'])}")
            print(f"      Text: {issue['text']}")
    else:
        print(f"\n⚠️  NO POTENTIAL ISSUES DETECTED!")
        print(f"\n   This suggests the document structure doesn't match what the parser expects.")
        print(f"   The parser looks for:")
        print(f"   1. Plain text bullets (•, -, *) at the start of lines")
        print(f"   2. Format: • Issue Title: Description...")

    # Show bullet format analysis
    print(f"\n📌 BULLET FORMAT ANALYSIS:")
    list_items = [p for p in potential_issues if p.get('is_list_item')]
    plain_bullets = [p for p in potential_issues if p.get('bullet_char')]

    print(f"  Paragraphs with Word list formatting: {len(list_items)}")
    print(f"  Paragraphs with plain text bullets: {len(plain_bullets)}")

    if len(list_items) > 0 and len(plain_bullets) == 0:
        print(f"\n  ⚠️  ISSUE DETECTED: Document uses Word's built-in bullet lists!")
        print(f"      The parser expects plain text bullets (• - *)")
        print(f"      But this document uses rich text bullet formatting.")
        print(f"\n  💡 SOLUTION: Parser needs to check paragraph._p.pPr.numPr for list items")

    if detailed and issues_list:
        print(f"\n📝 DETAILED ISSUE ANALYSIS (First 3):")
        for i, issue in enumerate(issues_list[:3], 1):
            print(f"\n  Issue #{i}:")
            print(f"  Text: {issue['text']}")
            print(f"  Full details: {issue}")

    return {
        'sections': len(sections_found),
        'potential_issues': len(issues_list),
        'uses_list_formatting': len(list_items) > 0,
        'uses_plain_bullets': len(plain_bullets) > 0
    }


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/diagnose_docx.py <path_to_audit.docx_or_directory>")
        sys.exit(1)

    path = sys.argv[1]
    detailed = '--detailed' in sys.argv

    if os.path.isfile(path):
        # Single file
        analyze_document(path, detailed=detailed)

    elif os.path.isdir(path):
        # Directory - analyze first 3 DOCX files
        docx_files = [f for f in os.listdir(path) if f.endswith('.docx') and not f.startswith('~')]

        if not docx_files:
            print(f"❌ No DOCX files found in {path}")
            sys.exit(1)

        print(f"\n📁 Found {len(docx_files)} DOCX files in {path}")
        print(f"   Analyzing first 3 files...\n")

        results = []
        for docx_file in docx_files[:3]:
            result = analyze_document(os.path.join(path, docx_file), detailed=detailed)
            results.append(result)

        # Overall summary
        print(f"\n{'='*70}")
        print(f"📊 OVERALL ANALYSIS (3 files)")
        print(f"{'='*70}")

        all_use_list = all(r['uses_list_formatting'] for r in results)
        all_use_plain = all(r['uses_plain_bullets'] for r in results)

        print(f"  All files use Word list formatting: {all_use_list}")
        print(f"  All files use plain text bullets: {all_use_plain}")

        if all_use_list and not all_use_plain:
            print(f"\n  🎯 ROOT CAUSE IDENTIFIED:")
            print(f"     All documents use Word's built-in bullet list formatting,")
            print(f"     but the parser only checks for plain text bullets (•, -, *).")
            print(f"\n  💡 FIX NEEDED:")
            print(f"     Update document_parser.py line 254 to also check for")
            print(f"     paragraph._p.pPr.numPr (Word list formatting metadata)")

    else:
        print(f"❌ Path not found: {path}")
        sys.exit(1)

    print(f"\n✅ Analysis complete!")


if __name__ == "__main__":
    main()

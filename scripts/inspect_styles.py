#!/usr/bin/env python3
"""
DOCX Style Inspector

Shows all paragraph styles used in a DOCX file to understand its structure.

Usage:
    python3 scripts/inspect_styles.py path/to/audit.docx
"""

import sys
import os
from docx import Document
from collections import Counter


def inspect_document_styles(docx_path: str):
    """Show all styles used in the document."""
    print(f"\n{'='*70}")
    print(f"📄 Inspecting: {os.path.basename(docx_path)}")
    print(f"{'='*70}\n")

    doc = Document(docx_path)

    # Collect all styles and their usage
    style_usage = Counter()
    style_examples = {}

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()

        if text:  # Only count non-empty paragraphs
            style_usage[style] += 1

            # Save first example of each style
            if style not in style_examples:
                style_examples[style] = text[:100]

    # Show styles sorted by usage
    print(f"📊 PARAGRAPH STYLES FOUND ({len(style_usage)} unique styles):\n")

    for style, count in style_usage.most_common():
        example = style_examples[style]
        print(f"  {style:30s} ({count:3d}x)  Example: {example}")

    # Look for potential section headers
    print(f"\n🔍 POTENTIAL SECTION HEADERS (non-Normal styles with < 5 uses):\n")

    for style, count in style_usage.items():
        if count < 5 and style != 'Normal':
            example = style_examples[style]
            print(f"  {style:30s} ({count:2d}x)  {example}")

    # Show first 30 paragraphs with their styles
    print(f"\n📝 FIRST 30 PARAGRAPHS (with styles):\n")

    for i, para in enumerate(doc.paragraphs[:30], 1):
        text = para.text.strip()
        if text:
            print(f"  [{i:2d}] {para.style.name:20s} | {text[:80]}")

    print(f"\n✅ Inspection complete!")


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/inspect_styles.py <path_to_audit.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]

    if not os.path.isfile(docx_path):
        print(f"❌ File not found: {docx_path}")
        sys.exit(1)

    inspect_document_styles(docx_path)


if __name__ == "__main__":
    main()

"""
Document Parser for CRO Analyzer

Parses DOCX audit files and extracts structured information:
- Client name
- Audit sections (Navigation, Hero, Product Page, etc.)
- Issues with descriptions and recommendations
- Reference websites

Designed to work with audit format observed in sample audits.
"""

import re
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.document import Document as DocumentType


class AuditSection:
    """Represents a section of an audit (e.g., "Home Page", "Navigation")."""

    def __init__(self, title: str, content: List[str], issues: List[Dict[str, str]]):
        self.title = title
        self.content = content  # Raw paragraph text
        self.issues = issues  # Parsed issue dictionaries

    def __repr__(self):
        return f"<AuditSection '{self.title}' with {len(self.issues)} issues>"


class AuditDocument:
    """Represents a parsed CRO audit document."""

    def __init__(
        self,
        client_name: str,
        sections: List[AuditSection],
        reference_websites: List[str],
        overview: str = ""
    ):
        self.client_name = client_name
        self.sections = sections
        self.reference_websites = reference_websites
        self.overview = overview

    def __repr__(self):
        return f"<AuditDocument for '{self.client_name}' with {len(self.sections)} sections>"

    def to_dict(self) -> Dict:
        """Convert audit to dictionary for ChromaDB storage."""
        return {
            'client_name': self.client_name,
            'overview': self.overview,
            'reference_websites': self.reference_websites,
            'sections': [
                {
                    'title': section.title,
                    'issues': section.issues
                }
                for section in self.sections
            ]
        }


class DocumentParser:
    """
    Parser for CRO audit DOCX files.

    Expected document structure (based on sample audits):
    1. Title: "{Client Name} CRO Audit Documentation"
    2. Index section listing major sections
    3. "Audit Overview" section
    4. "Reference Websites" section
    5. Individual sections with bulleted issues
    """

    def __init__(self, docx_path: str):
        """
        Initialize parser with DOCX file path.

        Args:
            docx_path: Path to DOCX file
        """
        self.docx_path = docx_path
        self.document = Document(docx_path)

    def parse(self) -> AuditDocument:
        """
        Parse the entire audit document.

        Returns:
            AuditDocument with extracted sections and issues
        """
        # Extract client name from title
        client_name = self._extract_client_name()

        # Extract overview
        overview = self._extract_overview()

        # Extract reference websites
        reference_websites = self._extract_reference_websites()

        # Extract all sections with issues
        sections = self._extract_sections()

        return AuditDocument(
            client_name=client_name,
            sections=sections,
            reference_websites=reference_websites,
            overview=overview
        )

    def _extract_client_name(self) -> str:
        """
        Extract client name from document title.

        Expected format: "{Client Name} CRO Audit Documentation"
        """
        for paragraph in self.document.paragraphs[:10]:  # Check first 10 paragraphs
            text = paragraph.text.strip()
            if "CRO Audit Documentation" in text:
                # Remove "CRO Audit Documentation" to get client name
                client_name = text.replace("CRO Audit Documentation", "").strip()
                print(f"✓ Detected client: {client_name}")
                return client_name

        # Fallback: use filename
        import os
        filename = os.path.basename(self.docx_path)
        client_name = filename.replace('.docx', '').replace('_', ' ')
        print(f"⚠ Could not find client name in document, using filename: {client_name}")
        return client_name

    def _extract_overview(self) -> str:
        """
        Extract the "Audit Overview" section content.

        Returns:
            Full text of the overview section
        """
        in_overview = False
        overview_lines = []

        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()

            # Start capturing when we hit "Audit Overview"
            if "Audit Overview" in text and paragraph.style.name.startswith('Heading'):
                in_overview = True
                continue

            # Stop when we hit another heading
            if in_overview and paragraph.style.name.startswith('Heading'):
                break

            if in_overview and text:
                overview_lines.append(text)

        return '\n'.join(overview_lines)

    def _extract_reference_websites(self) -> List[str]:
        """
        Extract list of reference websites.

        Expected format: Section titled "Reference Websites" with bullet points
        """
        in_references = False
        websites = []

        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()

            # Start capturing when we hit "Reference Websites"
            if "Reference Websites" in text and paragraph.style.name.startswith('Heading'):
                in_references = True
                continue

            # Stop when we hit another heading
            if in_references and paragraph.style.name.startswith('Heading'):
                break

            if in_references and text:
                # Clean bullet points and extract site name
                cleaned = text.lstrip('•-*').strip()
                if cleaned and not cleaned.startswith('('):  # Skip parenthetical notes
                    websites.append(cleaned)

        print(f"✓ Found {len(websites)} reference websites")
        return websites

    def _extract_index_sections(self) -> List[str]:
        """
        Extract section names from the Index.

        Returns:
            List of section names
        """
        in_index = False
        section_names = []
        skip_words = ['index:', 'audit overview:', 'reference websites']

        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            lower_text = text.lower()

            # Start capturing when we hit "Index:"
            if lower_text.startswith('index:'):
                in_index = True
                continue

            # Stop at "Audit Overview:" or "Reference Websites:"
            if in_index and (lower_text.startswith('audit overview:') or lower_text.startswith('reference websites')):
                break

            # Collect section names
            if in_index and text and len(text) > 3:
                # Skip if it's a skippable keyword
                if not any(skip in lower_text for skip in skip_words):
                    section_names.append(text)

        return section_names

    def _is_section_header(self, text: str, index_sections: List[str]) -> bool:
        """
        Determine if a paragraph is a section header.

        Checks if text matches any index section (with fuzzy matching).
        """
        if not text or len(text) < 3:
            return False

        # Exact match with index sections
        if text in index_sections:
            return True

        # Fuzzy match (handles slight variations)
        text_lower = text.lower()
        for section in index_sections:
            section_lower = section.lower()
            # Check if they're very similar
            if text_lower == section_lower:
                return True
            # Check if one contains the other (for variations)
            if len(text) > 10 and (text_lower in section_lower or section_lower in text_lower):
                return True

        return False

    def _extract_sections(self) -> List[AuditSection]:
        """
        Extract all audit sections with their issues.

        Uses the Index to identify section boundaries since documents
        don't use heading styles - everything is "normal" style.

        Returns:
            List of AuditSection objects
        """
        # First, extract section names from Index
        index_sections = self._extract_index_sections()
        print(f"  Found {len(index_sections)} sections in Index")

        sections = []
        current_section = None
        current_paragraphs = []  # Store Paragraph objects, not just text
        skip_sections = ['Index', 'Audit Overview', 'Reference Websites']

        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()

            # Check if this paragraph is a section header
            if self._is_section_header(text, index_sections):
                # Save previous section if exists
                if current_section and current_section not in skip_sections:
                    issues = self._parse_issues_from_paragraphs(current_paragraphs)
                    sections.append(AuditSection(
                        title=current_section,
                        content=[p.text.strip() for p in current_paragraphs],
                        issues=issues
                    ))

                # Start new section
                current_section = text
                current_paragraphs = []

            elif current_section and text:
                # Add paragraph to current section
                current_paragraphs.append(paragraph)

        # Save last section
        if current_section and current_section not in skip_sections:
            issues = self._parse_issues_from_paragraphs(current_paragraphs)
            sections.append(AuditSection(
                title=current_section,
                content=[p.text.strip() for p in current_paragraphs],
                issues=issues
            ))

        print(f"✓ Extracted {len(sections)} sections")
        return sections

    def _is_bullet_paragraph(self, paragraph: Paragraph) -> bool:
        """
        Check if a paragraph is a bullet list item.

        Checks both plain text bullets and Word list formatting.
        """
        text = paragraph.text.strip()

        # Method 1: Plain text bullet at start
        if re.match(r'^[•\-\*]\s+', text):
            return True

        # Method 2: Word list formatting (check paragraph properties)
        if paragraph._p.pPr is not None:
            numPr = paragraph._p.pPr.numPr
            if numPr is not None:
                return True

        # Method 3: Indented paragraph (might be a bullet without the character)
        # Check if paragraph has left indentation
        if paragraph.paragraph_format.left_indent and paragraph.paragraph_format.left_indent > 0:
            # If it's indented and short enough to be a title, likely an issue
            if len(text) < 200 and len(text) > 10:
                return True

        return False

    def _parse_issues_from_paragraphs(self, paragraphs: List[Paragraph]) -> List[Dict[str, str]]:
        """
        Parse issues from a list of Paragraph objects.

        Detects issues using multiple methods:
        1. Plain text bullets (•, -, *)
        2. Word list formatting
        3. Indented paragraphs
        4. Content patterns (colon separators, etc.)

        Returns:
            List of issue dictionaries with keys: title, description, why_it_matters, recommendations
        """
        issues = []
        current_issue = None

        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Check if this paragraph starts a new issue
            is_bullet = self._is_bullet_paragraph(paragraph)

            if is_bullet:
                # Save previous issue
                if current_issue:
                    issues.append(current_issue)

                # Start new issue
                # Remove bullet character if it's a plain text bullet
                cleaned_text = re.sub(r'^[•\-\*]\s+', '', text).strip()

                # Try to split title and description by colon
                if ':' in cleaned_text and cleaned_text.index(':') < 100:
                    title, description = cleaned_text.split(':', 1)
                    current_issue = {
                        'title': title.strip(),
                        'description': description.strip(),
                        'why_it_matters': '',
                        'recommendations': []
                    }
                else:
                    current_issue = {
                        'title': cleaned_text,
                        'description': '',
                        'why_it_matters': '',
                        'recommendations': []
                    }

            elif current_issue:
                # Continuation of current issue
                lower_text = text.lower()

                if 'why it matters' in lower_text or 'why this matters' in lower_text:
                    # Extract rationale
                    current_issue['why_it_matters'] += text + ' '

                elif 'recommendation' in lower_text or 'solution' in lower_text or 'fix' in lower_text:
                    # Extract recommendation
                    current_issue['recommendations'].append(text)

                else:
                    # Add to description
                    if current_issue['description']:
                        current_issue['description'] += ' ' + text
                    else:
                        current_issue['description'] = text

        # Save last issue
        if current_issue:
            issues.append(current_issue)

        # Clean up whitespace
        for issue in issues:
            issue['description'] = issue['description'].strip()
            issue['why_it_matters'] = issue['why_it_matters'].strip()
            issue['recommendations'] = [r.strip() for r in issue['recommendations'] if r.strip()]

        return issues

    def _parse_issues(self, content: List[str]) -> List[Dict[str, str]]:
        """
        Parse issues from section content (legacy method for backward compatibility).

        Expected pattern:
        • Issue Title: Issue description...
          Why it matters...
          Recommendation...

        Returns:
            List of issue dictionaries with keys: title, description, why_it_matters, recommendations
        """
        issues = []
        current_issue = None

        for line in content:
            # Detect new issue (starts with bullet: •, -, *)
            if re.match(r'^[•\-\*]\s+', line):
                # Save previous issue
                if current_issue:
                    issues.append(current_issue)

                # Start new issue
                cleaned_line = re.sub(r'^[•\-\*]\s+', '', line).strip()

                # Try to split title and description by colon
                if ':' in cleaned_line:
                    title, description = cleaned_line.split(':', 1)
                    current_issue = {
                        'title': title.strip(),
                        'description': description.strip(),
                        'why_it_matters': '',
                        'recommendations': []
                    }
                else:
                    current_issue = {
                        'title': cleaned_line,
                        'description': '',
                        'why_it_matters': '',
                        'recommendations': []
                    }

            elif current_issue:
                # Continuation of current issue
                lower_line = line.lower()

                if 'why it matters' in lower_line or 'why this matters' in lower_line:
                    # Extract rationale
                    current_issue['why_it_matters'] += line + ' '

                elif 'recommendation' in lower_line or 'solution' in lower_line or 'fix' in lower_line:
                    # Extract recommendation
                    current_issue['recommendations'].append(line)

                else:
                    # Add to description
                    if current_issue['description']:
                        current_issue['description'] += ' ' + line
                    else:
                        current_issue['description'] = line

        # Save last issue
        if current_issue:
            issues.append(current_issue)

        # Clean up whitespace
        for issue in issues:
            issue['description'] = issue['description'].strip()
            issue['why_it_matters'] = issue['why_it_matters'].strip()
            issue['recommendations'] = [r.strip() for r in issue['recommendations'] if r.strip()]

        return issues


# Usage example
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python3 document_parser.py <path_to_audit.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    parser = DocumentParser(docx_path)
    audit = parser.parse()

    print(f"\n📄 Audit Summary")
    print(f"Client: {audit.client_name}")
    print(f"Sections: {len(audit.sections)}")
    print(f"Reference Sites: {len(audit.reference_websites)}")

    print(f"\n📋 Sections:")
    for section in audit.sections:
        print(f"  - {section.title}: {len(section.issues)} issues")

    # Print first issue as example
    if audit.sections and audit.sections[0].issues:
        first_issue = audit.sections[0].issues[0]
        print(f"\n📌 Example Issue:")
        print(f"Title: {first_issue['title']}")
        print(f"Description: {first_issue['description'][:100]}...")
        print(f"Recommendations: {len(first_issue['recommendations'])}")
